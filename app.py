import streamlit as st
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
import os

st.set_page_config(page_title="XML to DOCX Highlighter", layout="wide")
st.title("XML to DOCX Batch Highlighter 📝➡️📄")

# Sidebar uploads and reset
st.sidebar.header("Controls")
ref_file = st.sidebar.file_uploader("Reference DOCX", type=["docx"])
xml_files = st.sidebar.file_uploader("XML Files", type=["xml"], accept_multiple_files=True)
if st.sidebar.button("Reset"):
    st.experimental_rerun()

if st.sidebar.button("Run Highlighting"):
    if not ref_file or not xml_files:
        st.sidebar.error("Please upload both reference DOCX and at least one XML file.")
    else:
        def extract_reference_strings(doc_bytes):
            doc = Document(BytesIO(doc_bytes))
            strings = set()
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = getattr(para.style, 'name', '').lower()
                has_num = para._p.find(qn('w:numPr')) is not None
                has_bold = any(run.bold for run in para.runs if run.text.strip())
                if has_bold or has_num or 'list paragraph' in style or 'bullet' in style or text.startswith(('•','-','*')):
                    cleaned = text.lstrip('•-–*0123456789. )').strip()
                    if cleaned:
                        strings.add(cleaned)
            return sorted(strings, key=len, reverse=True)

        ref_strings = extract_reference_strings(ref_file.read())
        results = []
        st.write("## Processing files…")
        for xml in xml_files:
            xml_bytes = xml.read().decode('utf-8')
            lines = xml_bytes.splitlines()
            doc = Document()
            body = doc._body._element
            for child in list(body): body.remove(child)
            for line in lines:
                para = doc.add_paragraph()
                idx = 0
                length = len(line)
                while idx < length:
                    next_pos = length
                    next_ref = None
                    for ref in ref_strings:
                        pos = line.find(ref, idx)
                        if pos != -1 and pos < next_pos:
                            next_pos, next_ref = pos, ref
                    if next_ref:
                        if next_pos > idx:
                            para.add_run(line[idx:next_pos])
                        run = para.add_run(next_ref)
                        run.font.highlight_color = 7
                        idx = next_pos + len(next_ref)
                    else:
                        para.add_run(line[idx:])
                        break
            for para in list(doc.paragraphs):
                txt = para.text.strip().lower()
                if txt.isdigit() or 'generated by python-docx' in txt:
                    para._element.getparent().remove(para._element)
            out_buf = BytesIO()
            doc.save(out_buf)
            out_buf.seek(0)
            filename = os.path.splitext(xml.name)[0] + '_highlighted.docx'
            results.append((filename, out_buf))

        st.success("Batch highlighting complete!")
        st.write("## Download Results")
        for name, buf in results:
            st.download_button(label=f"Download {name}", data=buf, file_name=name, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

st.sidebar.markdown("---")
st.sidebar.write("Built with Streamlit and python-docx.")
